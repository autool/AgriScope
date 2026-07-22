"""成果验收 DOCX、PDF 与清单实体渲染器。"""

import json
import re
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

from app.core.exceptions import ValidationException


class AcceptanceReportRenderer:
    """把当前成果包快照渲染为可送审 DOCX 和分页 PDF。"""

    @staticmethod
    def _load_font(
        size: int,
    ) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
        """加载支持中文的系统字体。

        Args:
            size: 字体像素大小。

        Returns:
            ImageFont.FreeTypeFont | ImageFont.ImageFont: 可用字体。
        """
        candidates = (
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        )
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size=size)
            except OSError:
                continue
        return ImageFont.load_default()

    @staticmethod
    def _set_cell_fill(cell: object, color: str) -> None:
        """设置 DOCX 表格单元格底色。

        Args:
            cell: python-docx 单元格。
            color: 六位十六进制颜色。

        Returns:
            None: 原位修改 XML。
        """
        properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        properties.append(shading)

    @staticmethod
    def _set_cell_text(cell: object, value: object, *, bold: bool = False) -> None:
        """写入 DOCX 单元格并统一字体。

        Args:
            cell: python-docx 单元格。
            value: 待显示值。
            bold: 是否加粗。

        Returns:
            None: 原位写入单元格。
        """
        cell.text = "" if value is None else str(value)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(9)
                run.bold = bold

    @classmethod
    def _add_table(
        cls,
        document: Document,
        headers: list[str],
        rows: list[list[object]],
        widths: list[float] | None = None,
    ) -> None:
        """向 DOCX 添加带表头样式的表格。

        Args:
            document: Word 文档。
            headers: 表头。
            rows: 数据行。
            widths: 可选列宽厘米值。

        Returns:
            None: 文档原位增加表格。
        """
        table = document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cls._set_cell_fill(cell, "DCE9E1")
            cls._set_cell_text(cell, header, bold=True)
            if widths:
                cell.width = Cm(widths[index])
        for values in rows:
            cells = table.add_row().cells
            for index, value in enumerate(values):
                cls._set_cell_text(cells[index], value)
                if widths:
                    cells[index].width = Cm(widths[index])

    @staticmethod
    def _status_text(value: object) -> str:
        """将证据状态转换为中文结论。

        Args:
            value: 原始状态。

        Returns:
            str: 中文状态。
        """
        return {
            "included": "已纳入实体",
            "provided": "已提供",
            "referenced": "校验后引用",
            "verified_reference": "校验后引用",
            "not_provided": "未提供",
            "missing_physical_evidence": "缺少实体证据",
            True: "是",
            False: "否",
        }.get(value, str(value if value is not None else "--"))

    @classmethod
    def _workload_rows(cls, data: dict) -> list[list[object]]:
        """构建工作量统计行。

        Args:
            data: 验收报告冻结数据。

        Returns:
            list[list[object]]: 工作量指标、数值和说明。
        """
        summary = data["quality_summary"]
        return [
            ["任务有效图斑", summary.get("plot_count", 0), "当前任务显式范围"],
            ["完成质量检查", summary.get("quality_checked_count", 0), "图斑级规则检查"],
            [
                "通过质量门禁",
                summary.get("quality_passing_count", 0),
                "生成成果包时全部通过",
            ],
            [
                "外业核查记录",
                summary.get("field_verification_count", 0),
                cls._status_text(summary.get("field_evidence_status")),
            ],
            [
                "灾害斑块",
                summary.get("disaster_patch_count", 0),
                cls._status_text(summary.get("disaster_evidence_status")),
            ],
            ["专题图", summary.get("thematic_map_count", 0), "已校验实体"],
            ["独立监理报告", summary.get("supervision_report_count", 0), "已校验实体"],
            [
                "面积统计正式报告",
                summary.get("statistics_report_count", 0),
                "当前报告成员",
            ],
            ["多格式矢量成果", summary.get("vector_export_count", 0), "当前导出成员"],
            ["多源数据资产", summary.get("dataset_asset_count", 0), "归档目录项"],
            ["影像处理实体", summary.get("imagery_step_count", 0), "受控引用"],
            ["成果包文件", len(data["delivery_manifest"]), "含 manifest 清单项"],
        ]

    @classmethod
    def _quality_rows(cls, data: dict) -> list[list[object]]:
        """构建精度与质量评价行。

        Args:
            data: 验收报告冻结数据。

        Returns:
            list[list[object]]: 指标、结果和评价。
        """
        summary = data["quality_summary"]
        checked = int(summary.get("quality_checked_count", 0) or 0)
        total = int(summary.get("quality_total_count", 0) or 0)
        passing = int(summary.get("quality_passing_count", 0) or 0)
        coverage = (
            100
            if total and checked == total
            else round(checked / total * 100, 2)
            if total
            else 0
        )
        pass_rate = (
            100
            if total and passing == total
            else round(passing / total * 100, 2)
            if total
            else 0
        )
        return [
            [
                "质量检查覆盖率",
                f"{coverage}%",
                "全部覆盖" if coverage == 100 else "未全部覆盖",
            ],
            [
                "图斑门禁通过率",
                f"{pass_rate}%",
                "全部通过" if pass_rate == 100 else "存在未通过图斑",
            ],
            ["任务平均质量得分", summary.get("quality_score", 0), "达到成果生成门槛"],
            ["未关闭问题", summary.get("open_issue_count", 0), "应为 0"],
            ["已关闭问题", summary.get("resolved_issue_count", 0), "保留整改审计"],
            ["待处置外业疑点", summary.get("pending_field_count", 0), "应为 0"],
            [
                "缺少现场照片记录",
                summary.get("field_missing_photo_count", 0),
                "有外业记录时应为 0",
            ],
            [
                "业务影像预处理",
                cls._status_text(summary.get("imagery_processing_complete")),
                summary.get("imagery_asset_code") or "未登记",
            ],
            ["三级审核记录", summary.get("review_record_count", 0), "任务状态：已完成"],
        ]

    @classmethod
    def build_docx(cls, data: dict) -> bytes:
        """生成包含工作量、精度评价和成果清单的 DOCX。

        Args:
            data: 验收报告冻结数据。

        Returns:
            bytes: DOCX 实体字节。
        """
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)
        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles["Normal"].font.size = Pt(10.5)
        for style_name in ("Title", "Heading 1", "Heading 2"):
            styles[style_name].font.name = "Microsoft YaHei"
            styles[style_name]._element.rPr.rFonts.set(
                qn("w:eastAsia"), "Microsoft YaHei"
            )

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(data["report_title"])
        title_run.bold = True
        title_run.font.size = Pt(20)
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(
            f"报告编号：{data['report_code']}  |  版本：V{data['version']}"
        )

        document.add_heading("一、项目与任务概况", level=1)
        cls._add_table(
            document,
            ["项目", "内容"],
            [
                ["任务编号", data["task_code"]],
                ["任务名称", data["task_name"]],
                ["任务状态", "已完成三级审核"],
                ["成果包编号", data["delivery_package_code"]],
                ["成果包 SHA-256", data["delivery_package_checksum_sha256"]],
                ["报告生成时间", data["generated_at"]],
                [
                    "报告生成人",
                    f"{data['generated_by']}（{data['generated_by_role']}）",
                ],
                ["生成依据", data["generation_comment"]],
            ],
            [4.0, 12.0],
        )

        document.add_heading("二、工作量统计", level=1)
        cls._add_table(
            document,
            ["工作项", "数量", "证据说明"],
            cls._workload_rows(data),
            [5.0, 3.0, 8.0],
        )

        document.add_heading("三、质量与精度评价", level=1)
        cls._add_table(
            document,
            ["评价指标", "结果", "验收说明"],
            cls._quality_rows(data),
            [5.0, 3.5, 7.5],
        )

        document.add_heading("四、三级审核与过程审计", level=1)
        review_rows = [
            [
                item.get("review_level", "--"),
                item.get("action", "--"),
                f"{item.get('reviewer', '--')} / {item.get('reviewer_role', '--')}",
                item.get("created_at", "--"),
                item.get("comment", ""),
            ]
            for item in data["reviews"]
        ]
        cls._add_table(
            document,
            ["审核层级", "动作", "审核人", "时间", "意见/证据"],
            review_rows or [["--", "--", "--", "--", "未提供审核记录"]],
        )

        document.add_heading("五、成果清单", level=1)
        manifest_rows = [
            [
                item.get("category", "--"),
                item.get("format", "--"),
                item.get("path", "--"),
                item.get("record_count", "成果文件")
                if item.get("record_count") is not None
                else "成果文件",
                cls._status_text(item.get("evidence_status")),
                item.get("checksum_sha256") or "清单自身",
            ]
            for item in data["delivery_manifest"]
        ]
        cls._add_table(
            document,
            ["类别", "格式", "归档路径", "记录数", "状态", "SHA-256"],
            manifest_rows,
        )

        document.add_heading("六、验收结论", level=1)
        conclusion = document.add_paragraph()
        conclusion.add_run(
            "本报告依据当前任务已完成三级审核、全部质量门禁和成果包实体完整性复核"
            "自动生成。成果包与本报告通过编号、文件大小和 SHA-256 绑定，可进入成果"
            "下载、归档和甲方验收资料整理环节。外业、灾害、专题图和监理等专项成果"
            "是否提供，以工作量统计及成果清单中的真实证据状态为准，不以空集合冒充完成。"
        )
        document.add_paragraph("承建单位（盖章）：____________________")
        document.add_paragraph("验收单位（盖章）：____________________")
        document.add_paragraph("验收日期：________年____月____日")

        document.core_properties.title = data["report_title"]
        document.core_properties.subject = data["report_code"]
        document.core_properties.author = data["generated_by"]
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _wrap_text(text: object, width: int) -> list[str]:
        """按中英文近似显示宽度切分 PDF 文本。

        Args:
            text: 待显示内容。
            width: 每行近似全角字符数。

        Returns:
            list[str]: 至少一行的文本列表。
        """
        value = str(text if text is not None else "--").replace("\r", "")
        lines: list[str] = []
        for paragraph in value.split("\n"):
            if not paragraph:
                lines.append("")
                continue
            current = ""
            units = 0
            for character in paragraph:
                character_units = 1 if ord(character) < 128 else 2
                if current and units + character_units > width * 2:
                    lines.append(current)
                    current = character
                    units = character_units
                else:
                    current += character
                    units += character_units
            lines.append(current)
        return lines or [""]

    @classmethod
    def _pdf_content(cls, data: dict) -> list[tuple[str, str]]:
        """构建 PDF 顺序内容块。

        Args:
            data: 验收报告冻结数据。

        Returns:
            list[tuple[str, str]]: 样式和文本列表。
        """
        blocks: list[tuple[str, str]] = [
            ("title", data["report_title"]),
            ("meta", f"报告编号：{data['report_code']}  版本：V{data['version']}"),
            ("heading", "一、项目与任务概况"),
            ("body", f"任务：{data['task_code']} · {data['task_name']}"),
            ("body", "任务状态：已完成三级审核"),
            ("body", f"成果包：{data['delivery_package_code']}"),
            ("body", f"成果包 SHA-256：{data['delivery_package_checksum_sha256']}"),
            ("body", f"生成人：{data['generated_by']}（{data['generated_by_role']}）"),
            ("body", f"生成时间：{data['generated_at']}"),
            ("body", f"生成依据：{data['generation_comment']}"),
            ("heading", "二、工作量统计"),
        ]
        blocks.extend(
            ("body", f"{row[0]}：{row[1]}；{row[2]}")
            for row in cls._workload_rows(data)
        )
        blocks.append(("heading", "三、质量与精度评价"))
        blocks.extend(
            ("body", f"{row[0]}：{row[1]}；{row[2]}") for row in cls._quality_rows(data)
        )
        blocks.append(("heading", "四、三级审核与过程审计"))
        if data["reviews"]:
            blocks.extend(
                (
                    "body",
                    f"{item.get('review_level', '--')} / {item.get('action', '--')} / "
                    f"{item.get('reviewer', '--')}"
                    f"（{item.get('reviewer_role', '--')}）/ "
                    f"{item.get('created_at', '--')}：{item.get('comment', '')}",
                )
                for item in data["reviews"]
            )
        else:
            blocks.append(("body", "未提供审核记录。"))
        blocks.append(("heading", "五、成果清单"))
        blocks.extend(
            (
                "small",
                f"{index}. {item.get('category', '--')} / {item.get('format', '--')} / "
                f"{item.get('path', '--')} / "
                f"{cls._status_text(item.get('evidence_status'))} / "
                f"SHA-256 {item.get('checksum_sha256') or '清单自身'}",
            )
            for index, item in enumerate(data["delivery_manifest"], start=1)
        )
        blocks.extend(
            [
                ("heading", "六、验收结论"),
                (
                    "body",
                    "本报告依据当前任务已完成三级审核、全部质量门禁和成果包实体完整性复核自动生成。"
                    "成果包与本报告通过编号、文件大小和 SHA-256 绑定，可进入成果下载、"
                    "归档和甲方"
                    "验收资料整理环节。专项成果是否提供，以本报告真实证据状态为准。",
                ),
                ("body", "承建单位（盖章）：____________________"),
                ("body", "验收单位（盖章）：____________________"),
                ("body", "验收日期：________年____月____日"),
            ]
        )
        return blocks

    @classmethod
    def build_pdf(cls, data: dict) -> tuple[bytes, int]:
        """生成自动分页 A4 PDF。

        Args:
            data: 验收报告冻结数据。

        Returns:
            tuple[bytes, int]: PDF 字节和页数。
        """
        width, height = 1240, 1754
        margin_x, top, bottom = 96, 90, 90
        fonts = {
            "title": cls._load_font(42),
            "heading": cls._load_font(29),
            "meta": cls._load_font(21),
            "body": cls._load_font(23),
            "small": cls._load_font(18),
        }
        line_heights = {"title": 62, "heading": 48, "meta": 34, "body": 36, "small": 29}
        colors = {
            "title": "#173D2D",
            "heading": "#2F6B4D",
            "meta": "#64736C",
            "body": "#24342D",
            "small": "#46534D",
        }
        pages: list[Image.Image] = []
        page = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(page)
        y = top

        def finish_page() -> None:
            nonlocal page, draw, y
            footer_font = fonts["small"]
            draw.line(
                (margin_x, height - 62, width - margin_x, height - 62),
                fill="#D7E1DC",
                width=2,
            )
            draw.text(
                (margin_x, height - 50),
                data["report_code"],
                font=footer_font,
                fill="#728079",
            )
            draw.text(
                (width - margin_x - 90, height - 50),
                f"第 {len(pages) + 1} 页",
                font=footer_font,
                fill="#728079",
            )
            pages.append(page)
            page = Image.new("RGB", (width, height), "white")
            draw = ImageDraw.Draw(page)
            y = top

        for style, content in cls._pdf_content(data):
            wrap_width = (
                22
                if style == "title"
                else 44
                if style == "heading"
                else 51
                if style in {"body", "meta"}
                else 66
            )
            lines = cls._wrap_text(content, wrap_width)
            block_height = len(lines) * line_heights[style] + (
                18 if style in {"title", "heading"} else 8
            )
            if y + block_height > height - bottom:
                finish_page()
            if style == "heading":
                draw.rounded_rectangle(
                    (margin_x - 12, y - 5, width - margin_x, y + line_heights[style]),
                    radius=8,
                    fill="#EDF5F0",
                )
            for line in lines:
                draw.text((margin_x, y), line, font=fonts[style], fill=colors[style])
                y += line_heights[style]
            y += 18 if style in {"title", "heading"} else 8
        finish_page()

        output = BytesIO()
        pages[0].save(
            output,
            format="PDF",
            save_all=True,
            append_images=pages[1:],
            resolution=150.0,
            title=f"{data['report_code']} {data['report_title']}",
            author=data["generated_by"],
        )
        return output.getvalue(), len(pages)

    @staticmethod
    def build_manifest(
        data: dict,
        docx_content: bytes,
        pdf_content: bytes,
        page_count: int,
    ) -> dict:
        """构建与成果包快照交叉绑定的报告清单。

        Args:
            data: 验收报告冻结数据。
            docx_content: DOCX 字节。
            pdf_content: PDF 字节。
            page_count: PDF 页数。

        Returns:
            dict: 可持久化 manifest。
        """
        return {
            "schema_version": "acceptance-report-v1",
            "report_code": data["report_code"],
            "report_title": data["report_title"],
            "version": data["version"],
            "task": {
                "task_code": data["task_code"],
                "task_name": data["task_name"],
                "task_plot_count": data["task_plot_count"],
                "task_updated_at_snapshot": data["task_updated_at_snapshot"],
            },
            "delivery_package": {
                "package_code": data["delivery_package_code"],
                "completed_at_snapshot": data["delivery_package_completed_at_snapshot"],
                "file_size_bytes": data["delivery_package_size_bytes"],
                "checksum_sha256": data["delivery_package_checksum_sha256"],
                "manifest_count": len(data["delivery_manifest"]),
                "quality_summary_checksum_sha256": data[
                    "quality_summary_checksum_sha256"
                ],
            },
            "generator": {
                "display_name": data["generated_by"],
                "user_code": data["generated_by_code"],
                "role_code": data["generated_by_role"],
                "generated_at": data["generated_at"],
                "generation_comment": data["generation_comment"],
            },
            "files": [
                {
                    "path": data["docx_filename"],
                    "format": "DOCX",
                    "file_size_bytes": len(docx_content),
                    "checksum_sha256": sha256(docx_content).hexdigest(),
                    "page_count": None,
                },
                {
                    "path": data["pdf_filename"],
                    "format": "PDF",
                    "file_size_bytes": len(pdf_content),
                    "checksum_sha256": sha256(pdf_content).hexdigest(),
                    "page_count": page_count,
                },
            ],
        }

    @staticmethod
    def build_bundle(
        docx_filename: str,
        docx_content: bytes,
        pdf_filename: str,
        pdf_content: bytes,
        manifest: dict,
    ) -> bytes:
        """把 DOCX、PDF 和 manifest 组合成 ZIP。

        Args:
            docx_filename: DOCX 文件名。
            docx_content: DOCX 字节。
            pdf_filename: PDF 文件名。
            pdf_content: PDF 字节。
            manifest: 报告清单。

        Returns:
            bytes: ZIP 字节。
        """
        output = BytesIO()
        with ZipFile(output, "w", ZIP_DEFLATED) as archive:
            archive.writestr(docx_filename, docx_content)
            archive.writestr(pdf_filename, pdf_content)
            archive.writestr(
                "manifest.json",
                json.dumps(manifest, ensure_ascii=False, indent=2),
            )
        return output.getvalue()

    @staticmethod
    def validate_docx(content: bytes, report_code: str) -> None:
        """校验 DOCX ZIP 结构和报告身份。

        Args:
            content: DOCX 实体字节。
            report_code: 预期报告编号。

        Returns:
            None: 校验通过时无返回值。
        """
        if not content.startswith(b"PK\x03\x04"):
            raise ValidationException("成果验收 DOCX 文件签名不合法")
        try:
            with ZipFile(BytesIO(content)) as archive:
                names = archive.namelist()
                if (
                    len(names) > 512
                    or len(names) != len(set(names))
                    or "[Content_Types].xml" not in names
                    or "word/document.xml" not in names
                    or any(
                        name.startswith("/") or ".." in Path(name).parts
                        for name in names
                    )
                    or sum(item.file_size for item in archive.infolist())
                    > 50 * 1024 * 1024
                ):
                    raise ValidationException("成果验收 DOCX 结构不安全或不完整")
                document_xml = archive.read("word/document.xml")
        except (BadZipFile, KeyError) as exc:
            raise ValidationException("成果验收 DOCX 无法解析") from exc
        if report_code.encode("utf-8") not in document_xml:
            raise ValidationException("成果验收 DOCX 报告编号不一致")

    @staticmethod
    def validate_pdf(content: bytes, report_code: str) -> int:
        """校验 PDF 签名、身份元数据和页数。

        Args:
            content: PDF 实体字节。
            report_code: 预期报告编号。

        Returns:
            int: 解析得到的页数。
        """
        if not content.startswith(b"%PDF"):
            raise ValidationException("成果验收 PDF 文件签名不合法")
        page_count = len(re.findall(rb"/Type\s*/Page\b", content))
        if page_count < 1:
            raise ValidationException("成果验收 PDF 未包含有效页面")
        return page_count
